"""
Preenchimento dos arquivos oficiais (Excel e Word) a partir dos dados do
processo e do resultado do calculo de tributos (calculo.ResultadoTributos).

Mapeamento de celulas do FORMULARIO_DSI__REMATE.xlsx (planilha "Plan1"),
levantado inspecionando o arquivo original celula a celula:

    B9   Nome / Nome Empresarial (importador)
    Q9   CPF/CNPJ
    B11  Endereco completo
    B13  RG / Passaporte
    Q13  Nacionalidade
    W13  Data do desembarque
    B19  Valor total dos bens (US$)
    H19  Valor do frete (US$)
    M19  Valor do seguro (US$)
    R19  Valor aduaneiro (R$)
    W19  Taxa de conversao (R$)
    P26  Pais de procedencia
    Y26  Termo de entrada
    B26  Transportador
    I26  Identificacao do veiculo
    U26  Data de chegada
    B28  No do conhecimento / etiqueta de bagagem
    L28  Qtde volumes
    P28  Peso bruto (kg)
    T28  Peso liquido (kg)
    X28  Depositario / armazem
    B31  Item (numero)
    D31  Qtde
    F31  Unidade
    I31  Descricao (linha 1) -- I32 e limpa (era a linha 2 do exemplo original)
    Z31  Valor FOB (US$) do item
    Z37  TOTAL (US$) -- so ha 1 item, entao repete o valor FOB
    B40  Local e data

A tabela "DEMONSTRATIVO DE CALCULO DOS TRIBUTOS" no Word tem 5 linhas.
A linha 3 (indice 3) tem os 14 valores do item; a linha 4 (indice 4, TOTAL)
so tem 4 celulas realmente independentes (impostos de cada tributo) -- o
resto da linha 4 e celula mesclada verticalmente com a linha 3 ou o texto
fixo "TOTAL", entao NAO deve ser escrito.
"""

from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass
from datetime import date

import docx
import openpyxl

from calculo import ResultadoTributos

# Texto fixo que sobrou de um preenchimento anterior dentro de uma caixa de
# texto no cabecalho do Word (fora do alcance do python-docx, que so ve
# tabelas/paragrafos normais). E especifico deste arquivo de template --
# se o template for trocado, este valor deixa de bater e simplesmente nao
# encontra nada para limpar (sem erro).
_NUMERO_DSI_RESIDUAL = "0917500-236227/2026"


@dataclass
class DadosProcesso:
    # Importador
    nome_importador: str = ""
    cpf_cnpj: str = ""
    endereco: str = ""
    rg_passaporte: str = ""
    nacionalidade: str = "BRASILEIRA"
    data_desembarque: date | None = None

    # Carga / item
    ncm: str = ""  # so digitos, sem pontos
    pais_procedencia: str = "PARAGUAI"
    termo_entrada: str = ""
    valor_frete_usd: float = 0.0
    valor_seguro_usd: float = 0.0
    taxa_conversao: float = 0.0
    item_qtde: int = 1
    item_unidade: str = "UNID."
    item_descricao: str = ""
    valor_fob_usd: float = 0.0
    local_data: str = ""

    # Opcionais -- raramente usados no dia a dia
    transportador: str = ""
    identificacao_veiculo: str = ""
    data_chegada: date | None = None
    numero_conhecimento: str = ""
    qtde_volumes: str = ""
    peso_bruto_kg: str = ""
    peso_liquido_kg: str = ""
    depositario: str = ""


def _fmt_moeda(valor: float) -> str:
    if not valor:
        return ""
    s = f"{valor:,.2f}"
    return s.replace(",", "§").replace(".", ",").replace("§", ".")


def _fmt_pct(valor: float) -> str:
    return f"{valor:.2f}".replace(".", ",")


def _fmt_data(d: date | None) -> str:
    return d.strftime("%d/%m/%Y") if d else ""


def preencher_excel(template_bytes: bytes, dados: DadosProcesso, r: ResultadoTributos) -> bytes:
    """Recebe os bytes do template .xlsx e devolve os bytes do arquivo preenchido."""
    wb = openpyxl.load_workbook(io.BytesIO(template_bytes))
    ws = wb["Plan1"] if "Plan1" in wb.sheetnames else wb.active

    ws["B9"] = dados.nome_importador
    ws["Q9"] = dados.cpf_cnpj
    ws["B11"] = dados.endereco
    ws["B13"] = dados.rg_passaporte
    ws["Q13"] = dados.nacionalidade
    ws["W13"] = _fmt_data(dados.data_desembarque)

    ws["B19"] = dados.valor_fob_usd or None
    ws["H19"] = dados.valor_frete_usd or None
    ws["M19"] = dados.valor_seguro_usd or None
    ws["R19"] = r.valor_mercadoria
    ws["W19"] = dados.taxa_conversao or None

    ws["P26"] = dados.pais_procedencia
    ws["Y26"] = dados.termo_entrada
    ws["B26"] = dados.transportador or None
    ws["I26"] = dados.identificacao_veiculo or None
    ws["U26"] = _fmt_data(dados.data_chegada) or None

    ws["B28"] = dados.numero_conhecimento or None
    ws["L28"] = dados.qtde_volumes or None
    ws["P28"] = dados.peso_bruto_kg or None
    ws["T28"] = dados.peso_liquido_kg or None
    ws["X28"] = dados.depositario or None

    ws["B31"] = 1
    ws["D31"] = dados.item_qtde
    ws["F31"] = dados.item_unidade
    ws["I31"] = dados.item_descricao
    ws["I32"] = None
    ws["Z31"] = dados.valor_fob_usd or None
    ws["Z37"] = dados.valor_fob_usd or None  # so 1 item -> total = valor do item

    ws["B40"] = dados.local_data

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _set_cell_text(cell, texto: str) -> None:
    """Atualiza o texto de uma celula de tabela do Word preservando a
    formatacao do run existente (fonte/tamanho), em vez de resetar o estilo
    como cell.text = ... faria."""
    if cell.paragraphs and cell.paragraphs[0].runs:
        cell.paragraphs[0].runs[0].text = texto
        for run in cell.paragraphs[0].runs[1:]:
            run.text = ""
    else:
        cell.text = texto


def _limpar_numero_dsi_residual(docx_bytes: bytes) -> bytes:
    """Remove o numero de DSI de um preenchimento anterior que ficou preso
    numa caixa de texto (o python-docx nao alcanca esse conteudo). Se o
    texto nao for encontrado (ex.: template trocado), devolve os bytes
    sem alterar nada."""
    entrada = zipfile.ZipFile(io.BytesIO(docx_bytes), "r")
    saida_buf = io.BytesIO()
    with zipfile.ZipFile(saida_buf, "w", zipfile.ZIP_DEFLATED) as saida:
        for item in entrada.infolist():
            conteudo = entrada.read(item.filename)
            if item.filename == "word/document.xml":
                conteudo = conteudo.replace(
                    _NUMERO_DSI_RESIDUAL.encode("utf-8"), b""
                )
            saida.writestr(item, conteudo)
    entrada.close()
    return saida_buf.getvalue()


def _encontrar_tabela_tributos(doc):
    """Acha a tabela "DEMONSTRATIVO DE CALCULO DOS TRIBUTOS" procurando pelo
    texto "CLASSIFICACAO FISCAL", em vez de assumir que ela e sempre
    doc.tables[0] -- programas diferentes convertem o .doc original para
    .docx de formas diferentes (o cabecalho "DSI No ___" vira uma tabela
    extra em alguns conversores e uma caixa de texto solta em outros),
    entao a posicao da tabela de tributos pode variar."""
    for tabela in doc.tables:
        for row in tabela.rows:
            for cell in row.cells:
                if "CLASSIFICAÇÃO FISCAL" in cell.text.upper():
                    return tabela
    return None


def preencher_word(template_bytes: bytes, dados: DadosProcesso, r: ResultadoTributos) -> bytes:
    """Recebe os bytes do template .docx (Demonstrativo de Calculo) e devolve
    os bytes do arquivo preenchido."""
    doc = docx.Document(io.BytesIO(template_bytes))

    tabela = _encontrar_tabela_tributos(doc)
    if tabela is None:
        raise ValueError(
            "Não encontrei a tabela \"DEMONSTRATIVO DE CÁLCULO DOS TRIBUTOS\" "
            "(procurei o texto 'CLASSIFICAÇÃO FISCAL' em todas as tabelas do "
            "arquivo). Confirme que templates/FORMULARIO_DSI_REMATE.docx é "
            "mesmo o Demonstrativo de Cálculo convertido do original."
        )
    if len(tabela.rows) < 5:
        raise ValueError(
            f"A tabela de tributos encontrada só tem {len(tabela.rows)} linha(s); "
            "esperava 5 (cabeçalho + linha do item + linha TOTAL). O arquivo "
            "pode estar corrompido ou ter sido editado."
        )

    linha_item = tabela.rows[3].cells
    valores = {
        0: "001",
        1: dados.ncm,
        2: _fmt_moeda(r.valor_mercadoria),
        3: _fmt_pct(r.aliquota_ii),
        4: _fmt_moeda(r.ii),
        5: _fmt_moeda(r.base_ipi),
        6: _fmt_pct(r.aliquota_ipi),
        7: _fmt_moeda(r.ipi),
        8: _fmt_moeda(r.valor_mercadoria),
        9: _fmt_pct(r.aliquota_pis),
        10: _fmt_moeda(r.pis),
        11: _fmt_moeda(r.valor_mercadoria),
        12: _fmt_pct(r.aliquota_cofins),
        13: _fmt_moeda(r.cofins),
    }
    for indice, texto in valores.items():
        _set_cell_text(linha_item[indice], texto)

    # Linha TOTAL -- so as 4 celulas de imposto sao independentes; o resto
    # e celula mesclada (com a linha do item ou com o rotulo "TOTAL") e nao
    # deve ser tocado.
    linha_total = tabela.rows[4].cells
    _set_cell_text(linha_total[4], _fmt_moeda(r.ii))
    _set_cell_text(linha_total[7], _fmt_moeda(r.ipi))
    _set_cell_text(linha_total[10], _fmt_moeda(r.pis))
    _set_cell_text(linha_total[13], _fmt_moeda(r.cofins))

    buf = io.BytesIO()
    doc.save(buf)
    return _limpar_numero_dsi_residual(buf.getvalue())